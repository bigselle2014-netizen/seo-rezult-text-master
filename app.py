import os
import re
import requests
import numpy as np
import streamlit as st
from datetime import datetime
from urllib.parse import quote
from docx import Document
from supabase import create_client, Client
from keep_alive import keep_alive

# =========================
# 🔐 КЛЮЧИ И КЛИЕНТЫ (ENV)
# =========================
SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")
PPLX_API_KEY = os.environ.get("PPLX_API_KEY")
TEXT_RU_KEY  = os.environ.get("TEXT_RU_KEY")

missing = [k for k, v in {
    "SUPABASE_URL": SUPABASE_URL,
    "SUPABASE_KEY": SUPABASE_KEY,
    "PPLX_API_KEY": PPLX_API_KEY,
    "TEXT_RU_KEY": TEXT_RU_KEY,
}.items() if not v]
if missing:
    st.error("❌ Не найдены переменные окружения: " + ", ".join(missing))
    st.stop()

supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

# =========================
# ⚙️ НАСТРОЙКИ UI
# =========================
st.set_page_config(page_title="SEO Rezult Text Master v7.4", layout="wide")

# =========================
# 👤 АВТОРИЗАЦИЯ (SIDEBAR)
# =========================
if "user" not in st.session_state:
    st.session_state.user = None

with st.sidebar:
    st.header("🔐 Авторизация")

    if st.session_state.user:
        st.success(f"Добро пожаловать, {st.session_state.user['email']}!")
        if st.button("Выйти"):
            st.session_state.user = None
            st.rerun()
    else:
        mode = st.radio("Действие", ["Вход", "Регистрация"], horizontal=True)
        email = st.text_input("Email")
        password = st.text_input("Пароль", type="password")
        if st.button("Продолжить"):
            try:
                if mode == "Регистрация":
                    if email.strip().lower() == "admin@seo-rezult.ru":
                        st.error("Регистрация под этим адресом запрещена.")
                    else:
                        res = supabase.auth.sign_up({"email": email, "password": password})
                        if getattr(res, "user", None):
                            st.success("✅ Регистрация успешна. Теперь войдите.")
                        else:
                            st.warning("Не удалось зарегистрировать. Проверьте email/пароль.")
                else:
                    res = supabase.auth.sign_in_with_password({"email": email, "password": password})
                    if getattr(res, "user", None):
                        st.session_state.user = {"email": email, "id": res.user.id}
                        st.rerun()
                    else:
                        st.error("Неверный email или пароль.")
            except Exception as e:
                msg = str(e)
                if "user_already_exists" in msg:
                    st.warning("⚠️ Такой email уже зарегистрирован. Войдите.")
                elif "invalid_credentials" in msg:
                    st.error("❌ Неверный email или пароль.")
                else:
                    st.error(f"Ошибка: {msg}")

# =========================
# 🧠 ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =========================
def perplexity_generate(prompt: str) -> str:
    headers = {"Authorization": f"Bearer {PPLX_API_KEY}", "Content-Type": "application/json"}
    payload = {
        "model": "sonar-pro",
        "messages": [{"role": "user", "content": [{"type": "text", "text": prompt}]}],
        "max_output_tokens": 2000
    }
    r = requests.post("https://api.perplexity.ai/chat/completions", json=payload, headers=headers, timeout=90)
    if not r.ok:
        st.error(f"Perplexity API ошибка {r.status_code}: {r.text}")
        r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]

def build_prompt(topic, site, competitors, lsi, banned, keys, symbols) -> str:
    return f"""
Ты опытный SEO-копирайтер и журналист.
Напиши экспертный SEO-текст на тему: {topic}
Сайт клиента: {site}
Конкуренты: {competitors}
Ключевые слова: {keys}
LSI-фразы (обязательно учесть): {lsi}
Запрещённые слова: {banned}
Объём ≈ {symbols} символов.
Пиши живым языком, без клише и “ИИ-тона”. Добавляй микро-примеры, детали и фактуру.
"""

def clean_text(text: str) -> str:
    return re.sub(r"[#*_>`]+", "", text).strip()

def check_missing_lsi(text: str, lsi_list: list[str]) -> list[str]:
    low = text.lower()
    return [w for w in lsi_list if w and w.lower() not in low]

def seo_score(text: str, keywords: str) -> dict:
    words = re.findall(r"\w+", text.lower())
    word_count = len(words) or 1
    avg_len = sum(len(w) for w in words) / word_count
    key_density = sum(text.lower().count(k.strip().lower())
                      for k in keywords.split(",") if k.strip()) / word_count * 100
    sentences = [s for s in re.split(r"[.!?]", text) if s.strip()]
    avg_sentence_len = (sum(len(s.split()) for s in sentences) / len(sentences)) if sentences else 0
    water = len(re.findall(r"\b(очень|это|также|поэтому|например|в целом|следовательно)\b", text.lower())) / word_count * 100
    return {
        "Количество слов": word_count,
        "Средняя длина слова": round(avg_len, 2),
        "Средняя длина предложения": round(avg_sentence_len, 2),
        "Плотность ключей (%)": round(key_density, 2),
        "Водность (%)": round(water, 2),
    }

def analyze_humanness(text: str) -> dict:
    sentences = [s.strip() for s in re.split(r"[.!?]", text) if s.strip()]
    words = re.findall(r"\w+", text.lower())
    unique_words = len(set(words)) or 1
    perplexity = round(np.exp(len(words) / unique_words), 2)
    lens = [len(s.split()) for s in sentences] or [0]
    burstiness = round(np.std(lens) / (np.mean(lens) + 1e-5), 2)
    bigrams = [f"{words[i]} {words[i+1]}" for i in range(max(0, len(words)-1))]
    repeats = len(bigrams) - len(set(bigrams))
    repeat_ratio = round(repeats / (len(bigrams) or 1) * 100, 2)
    human_score = 100 - ((perplexity / 50) * 20 + (repeat_ratio / 2) - (burstiness * 10))
    human_score = max(0, min(100, round(human_score, 1)))
    return {
        "Перплексия": perplexity,
        "Разнообразие предложений": burstiness,
        "Повторяемость фраз (%)": repeat_ratio,
        "Естественность текста (%)": human_score,
    }

def export_docx(text: str, report: dict, human_report: dict, filename="seo_text.docx"):
    doc = Document()
    doc.add_heading("SEO Rezult Text Master — Отчёт", level=1)
    doc.add_paragraph(text)
    doc.add_page_break()
    doc.add_heading("📊 SEO-анализ", level=2)
    for k, v in report.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.add_heading("🧠 Анализ естественности", level=2)
    for k, v in human_report.items():
        doc.add_paragraph(f"{k}: {v}")
    doc.save(filename)
    with open(filename, "rb") as f:
        st.download_button("📥 Скачать DOCX-отчёт", f, file_name=filename)

# =========================
# 🧭 ОСНОВНОЙ UI
# =========================
if not st.session_state.user:
    st.info("🔑 Войдите или зарегистрируйтесь, чтобы использовать генератор.")
    st.stop()

email = st.session_state.user["email"]
is_admin = (email.lower() == "admin@seo-rezult.ru")

st.title("🚀 SEO Rezult Text Master v7.4")
st.caption("Генератор SEO-текстов с LSI-итерациями, проверкой уникальности и естественности")

tab_labels = ["📝 Генерация", "📂 Мои тексты"]
if is_admin:
    tab_labels.append("🧑‍💼 Админ-панель")
tabs = st.tabs(tab_labels)

# -------------------------
# 📝 Генерация
# -------------------------
with tabs[0]:
    with st.form("input_form"):
        topic = st.text_input("Тематика текста")
        site = st.text_input("Сайт клиента")
        competitors = st.text_area("Ссылки на конкурентов (по одной в строке)")
        lsi_words = st.text_area("Список LSI-слов (через запятую)")
        banned = st.text_area("Запрещённые слова (через запятую)")
        keywords = st.text_area("Ключевые слова (через запятую)")
        symbols = st.number_input("Количество символов", value=8000, step=500)
        submitted = st.form_submit_button("Сгенерировать текст")

    if submitted:
        lsi_list = [w.strip() for w in lsi_words.split(",") if w.strip()]
        st.info("⚙️ Этап 1: Генерация текста...")
        base = build_prompt(topic, site, competitors, lsi_words, banned, keywords, symbols)
        text = clean_text(perplexity_generate(base))

        iteration = 1
        progress = st.progress(0)
        while True:
            missing = check_missing_lsi(text, lsi_list)
            if not missing:
                break
            st.warning(f"Этап 2: добавляем {len(missing)} LSI-слов(а)...")
            fix_prompt = (
                f"Добавь в текст слова {', '.join(missing)}.\n"
                f"Исходный текст:\n{text}"
            )
            addition = clean_text(perplexity_generate(fix_prompt))
            text += "\n" + addition
            iteration += 1
            progress.progress(min(90, iteration * 20))
        progress.progress(100)

        st.success("✅ Текст готов!")
        st.text_area("Результат", text, height=400)

        st.info("🔎 Проверка уникальности (Text.ru)…")
        try:
            r = requests.post("https://api.text.ru/post", data={"text": text, "userkey": TEXT_RU_KEY}, timeout=30)
            if r.ok and "text_uid" in r.json():
                uid = r.json()["text_uid"]
                res = requests.get("https://api.text.ru/post", params={"uid": uid, "userkey": TEXT_RU_KEY}, timeout=30).json()
                uniq = res.get("text_unique", "?")
                st.write(f"**Уникальность:** {uniq}%")
        except Exception as e:
            st.warning(f"Text.ru недоступен: {e}")

        st.info("📊 SEO-анализ")
        report = seo_score(text, keywords)
        st.table(report.items())

        st.info("🧠 Анализ естественности")
        human_report = analyze_humanness(text)
        st.table(human_report.items())

        export_docx(text, report, human_report)

        try:
            supabase.table("history").insert({
                "user_id": st.session_state.user["id"],
                "email": email,
                "date": datetime.now().isoformat(),
                "topic": topic,
                "symbols": int(symbols),
                "lsi_count": len(lsi_list),
                "text": text
            }).execute()
        except Exception as e:
            st.warning(f"Не удалось сохранить историю: {e}")

        st.markdown(f"[🧩 Проверить на AI Detector](https://aidetectorwriter.com/ru/?text={quote(text)})")

# -------------------------
# 📂 Мои тексты
# -------------------------
with tabs[1]:
    st.subheader("📂 Мои тексты")
    try:
        data = supabase.table("history").select("*").eq("user_id", st.session_state.user["id"]).order("date", desc=True).execute()
        rows = data.data or []
        if not rows:
            st.info("Пока нет сохранённых текстов.")
        for row in rows:
            with st.expander(f"{row.get('topic','—')} — {row.get('date','')}"):
                st.write(row.get("text","")[:600] + "…")
                st.caption(f"Символов: {row.get('symbols','?')}, LSI: {row.get('lsi_count','?')}")
                if st.button(f"🗑 Удалить", key=f"del_{row['id']}"):
                    supabase.table("history").delete().eq("id", row["id"]).execute()
                    st.rerun()
    except Exception as e:
        st.warning(f"Ошибка загрузки истории: {e}")

# -------------------------
# 🧑‍💼 Админ-панель
# -------------------------
if is_admin and len(tabs) > 2:
    with tabs[2]:
        st.subheader("🧑‍💼 Админ-панель (все тексты)")
        try:
            data = supabase.table("history").select("*").order("date", desc=True).execute()
            rows = data.data or []
            if not rows:
                st.info("База пуста.")
            for row in rows:
                with st.expander(f"{row.get('email','?')} — {row.get('topic','—')} — {row.get('date','')}"):
                    st.write(row.get("text","")[:600] + "…")
                    st.caption(f"Символов: {row.get('symbols','?')}, LSI: {row.get('lsi_count','?')}")
                    if st.button(f"🗑 Удалить {row['id']}", key=f"adm_del_{row['id']}"):
                        supabase.table("history").delete().eq("id", row["id"]).execute()
                        st.rerun()
        except Exception as e:
            st.warning(f"Ошибка админ-панели: {e}")

# =========================
# 🚀 Keep Alive
# =========================
keep_alive()
